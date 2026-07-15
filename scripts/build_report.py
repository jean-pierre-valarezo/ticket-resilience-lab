from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as PdfImage,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
EVIDENCE = DOCS / "evidence"
LOGO = DOCS / "ups-logo-color.png"
DIAGRAM = DOCS / "architecture-distribution.png"
DOCX_OUT = DOCS / "Informe_Ticket_Resilience_Lab.docx"
PDF_OUT = DOCS / "Informe_Ticket_Resilience_Lab.pdf"
MD_OUT = DOCS / "informe-tecnico.md"

TITLE = "Ticket Resilience Lab"
SUBTITLE = "Práctica: mecanismos de tolerancia a fallas en Kubernetes"
AUTHORS = "Alexander Chuquipoma y Jean Pierre Valarezo"
CAREER = "Carrera de Ingeniería en Ciencias de la Computación"
COURSE = "Sistemas Distribuidos"
TEACHER = "Ing. Cristian Timbi"
DATE = "14/07/2026"
LOCATION = "Cuenca - Ecuador"
REPO_URL = "https://github.com/jean-pierre-valarezo/ticket-resilience-lab.git"


@dataclass
class Evidence:
    caption: str
    filename: str


COMPONENTS = [
    ["API Gateway", "Entrada del sistema", "Recibe solicitudes de clientes y aplica rate limiting.", "gateway"],
    ["Reservas Core", "Orquestacion", "Coordina inventario, pago y notificación.", "reservations"],
    ["Inventario", "Consistencia crítica", "Valida y reserva asientos con PostgreSQL.", "inventory"],
    ["Pagos", "Stub externo", "Simula latencia y fallos realistas de una pasarela.", "payments"],
    ["Notificaciones", "Stub no crítico", "Simula envio de confirmaciones por correo.", "notifications"],
    ["Base de datos", "Persistencia", "Almacena estado de asientos y soporta bloqueo transaccional.", "postgres"],
]

FAILURE_MATRIX = [
    [
        "Inventario Fantasma",
        "Disponibilidad",
        "Escalar deployment/inventory a 0 réplicas",
        "Retry con backoff y fallo controlado",
        "Implementado",
    ],
    [
        "Pasarela Lenta",
        "Latencia",
        "PAYMENT_DELAY_MS=20000",
        "Timeout + circuit breaker",
        "Implementado",
    ],
    [
        "Diluvio de Peticiones",
        "Sobrecarga",
        "25 solicitudes consecutivas contra POST /reserve",
        "Rate limiting en gateway",
        "Implementado",
    ],
    [
        "Base de Datos Intermitente",
        "Conectividad",
        "Flapping de conexion o reinicio controlado de PostgreSQL",
        "Failover, backoff, idempotencia y outbox",
        "Análisis teórico",
    ],
    [
        "Correo Perdido",
        "Fallo no crítico",
        "Escalar deployment/notifications a 0 réplicas",
        "Fallback con notification.status=pending",
        "Implementado",
    ],
    [
        "Condición de Carrera",
        "Consistencia",
        "Solicitudes concurrentes sobre el mismo asiento",
        "SELECT FOR UPDATE, holds e idempotencia",
        "Análisis teórico",
    ],
]

IMPLEMENTED_FAILURES = [
    {
        "title": "Inventario Fantasma",
        "risk": "Sin inventario, el sistema podria confirmar una reserva sin comprobar disponibilidad real.",
        "command": ".\\scripts\\05-chaos-inventory-down.ps1",
        "defense": (
            "El servicio reservations aplica retry con backoff al invocar inventory. Si la dependencia no responde, "
            "se devuelve un error controlado y no se continua con pago ni notificación."
        ),
        "why": (
            "Retry con backoff es adecuado porque una caida breve de pods o endpoints puede ser transitoria. "
            "No se usa fallback positivo porque inventario es una dependencia crítica: inventar disponibilidad "
            "seria técnicamente incorrecto."
        ),
        "result": "La evidencia muestra HTTP 503. El sistema falla rapido y no confirma una reserva fantasma.",
        "evidence": "09-chaos-inventory-down-http-503.png",
    },
    {
        "title": "Pasarela Lenta",
        "risk": "Una pasarela de pagos lenta puede dejar hilos/conexiones esperando y propagar latencia al usuario.",
        "command": ".\\scripts\\07-chaos-payments-slow.ps1",
        "defense": (
            "reservations usa timeout al llamar payments y mantiene un circuit breaker en memoria. Tras tres fallos, "
            "el circuito pasa a open y rechaza nuevas llamadas sin esperar los 20 segundos de latencia."
        ),
        "why": (
            "El timeout limita la espera por solicitud y el circuit breaker evita seguir llamando una dependencia "
            "degradada. Unicamente reintentar seria peligroso porque amplificaria la carga sobre payments."
        ),
        "result": "Los intentos devuelven HTTP 503 y /resilience reporta payment_circuit.state=open.",
        "evidence": "12-chaos-payments-slow-circuit-breaker-open-2.png",
    },
    {
        "title": "Correo Perdido",
        "risk": "Una falla de notificaciónes podria cancelar ventas ya pagadas si se trata como dependencia crítica.",
        "command": ".\\scripts\\09-chaos-notifications-down.ps1",
        "defense": (
            "notifications se considera dependencia no crítica. Si no responde, reservations conserva la reserva "
            "confirmada y marca la notificación como pending para reintento operativo posterior."
        ),
        "why": (
            "Fallback es el patrón correcto porque el correo no define la validez de la entrada. Cancelar la reserva "
            "por no enviar email degradaria la disponibilidad sin proteger consistencia."
        ),
        "result": "La evidencia muestra HTTP 200, status=confirmed y notification.status=pending.",
        "evidence": "15-chaos-notifications-down-fallback-pending.png",
    },
    {
        "title": "Diluvio de Peticiones",
        "risk": "Un pico de tráfico puede saturar el gateway y arrastrar servicios internos.",
        "command": ".\\scripts\\11-chaos-traffic-spike.ps1",
        "defense": (
            "gateway mantiene un contador temporal por cliente y corta el exceso de solicitudes con HTTP 429. "
            "Los conflictos funcionales se reportan como HTTP 409, separados de la sobrecarga."
        ),
        "why": (
            "Rate limiting debe ubicarse en el borde para rechazar pronto el exceso de tráfico. Un circuit breaker "
            "interno no evitaria que las solicitudes entren al sistema."
        ),
        "result": "La evidencia muestra respuestas HTTP 429 despues de superar el limite configurado.",
        "evidence": "17-chaos-traffic-spike-rate-limit-429.png",
    },
]

EVIDENCE_LIST = [
    Evidence("Cluster Kubernetes con dos nodos en estado Ready.", "01-k8s-dos-nodos-ready.png"),
    Evidence("Pods desplegados: inventory aparece replicado entre ticket-lab y ticket-lab-m02.", "02-k8s-pods-inventory-replicado-dos-nodos.png"),
    Evidence("Port-forward hacia gateway, reservations e inventory.", "03-port-forward-servicios-locales.png"),
    Evidence("Health checks e inventario inicial.", "04-baseline-health-inventario-inicial-1.png"),
    Evidence("Reserva base confirmada.", "06-baseline-reserva-exitosa.png"),
    Evidence("Inventario posterior con asiento reservado.", "07-baseline-inventario-despues-1.png"),
    Evidence("Inventory down con HTTP 503.", "09-chaos-inventory-down-http-503.png"),
    Evidence("Payments lento con circuit breaker abierto.", "12-chaos-payments-slow-circuit-breaker-open-2.png"),
    Evidence("Notifications caído con fallback pending.", "15-chaos-notifications-down-fallback-pending.png"),
    Evidence("Traffic spike con HTTP 429.", "17-chaos-traffic-spike-rate-limit-429.png"),
]

REFERENCES = [
    "Brewer, E. A. (2012). CAP twelve years later: How the rules have changed. Computer, 45(2), 23-29. https://doi.org/10.1109/MC.2012.37",
    "Fowler, M. (2014). Circuit Breaker. https://martinfowler.com/bliki/CircuitBreaker.html",
    "Kubernetes. (2026). Deployments. https://kubernetes.io/docs/concepts/workloads/controllers/deployment/",
    "Kubernetes. (2026). Services, Load Balancing, and Networking. https://kubernetes.io/docs/concepts/services-networking/",
    "Kubernetes. (2026). Configure Liveness, Readiness and Startup Probes. https://kubernetes.io/docs/tasks/configure-pod-container/configure-liveness-readiness-startup-probes/",
    "Kubernetes. (2026). Pod Topology Spread Constraints. https://kubernetes.io/docs/concepts/scheduling-eviction/topology-spread-constraints/",
    "Kubernetes. (2026). Horizontal Pod Autoscaling. https://kubernetes.io/docs/tasks/run-application/horizontal-pod-autoscale/",
    "Nygard, M. T. (2018). Release It! Design and Deploy Production-Ready Software (2nd ed.). Pragmatic Bookshelf.",
    "PostgreSQL Global Development Group. (2026). Explicit Locking. https://www.postgresql.org/docs/current/explicit-locking.html",
    "PostgreSQL Global Development Group. (2026). SELECT. https://www.postgresql.org/docs/current/sql-select.html",
]


def arial(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    paths = [
        Path("C:/Windows/Fonts/arialbd.ttf") if bold else Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf") if bold else Path("C:/Windows/Fonts/calibri.ttf"),
    ]
    for path in paths:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def box(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], text: str, fill: str = "#ffffff") -> None:
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=14, fill=fill, outline="#111111", width=3)
    f = arial(24, True)
    lines: list[str] = []
    for piece in text.split("\n"):
        lines.extend(wrap(piece, 18) or [""])
    y = y1 + ((y2 - y1) - len(lines) * 30) // 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=f)
        draw.text((x1 + (x2 - x1 - (bbox[2] - bbox[0])) // 2, y), line, fill="#111111", font=f)
        y += 30


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int]) -> None:
    draw.line([start, end], fill="#111111", width=4)
    ex, ey = end
    sx, _ = start
    if ex >= sx:
        points = [(ex, ey), (ex - 15, ey - 9), (ex - 15, ey + 9)]
    else:
        points = [(ex, ey), (ex + 15, ey - 9), (ex + 15, ey + 9)]
    draw.polygon(points, fill="#111111")


def build_diagram() -> None:
    img = Image.new("RGB", (1600, 900), "#ffffff")
    draw = ImageDraw.Draw(img)
    title_font = arial(34, True)
    small = arial(20)
    draw.text((55, 35), "Arquitectura y distribucion multi-nodo", fill="#111111", font=title_font)
    draw.rounded_rectangle((55, 110, 760, 820), radius=24, outline="#111111", width=4, fill="#f3f3f3")
    draw.rounded_rectangle((840, 110, 1545, 820), radius=24, outline="#111111", width=4, fill="#f8f8f8")
    draw.text((95, 135), "Nodo ticket-lab", fill="#111111", font=arial(28, True))
    draw.text((880, 135), "Nodo ticket-lab-m02", fill="#111111", font=arial(28, True))
    box(draw, (80, 435, 300, 535), "Cliente\nDemo")
    box(draw, (170, 250, 420, 360), "Reservations\nCore")
    box(draw, (470, 250, 720, 360), "Inventory\nreplica A")
    box(draw, (950, 210, 1200, 320), "API\nGateway")
    box(draw, (950, 380, 1200, 490), "Payments\nStub")
    box(draw, (950, 550, 1200, 660), "Notifications\nStub")
    box(draw, (1245, 210, 1495, 320), "Inventory\nreplica B")
    box(draw, (1245, 420, 1495, 530), "PostgreSQL")
    arrow(draw, (300, 485), (950, 265))
    arrow(draw, (950, 265), (420, 305))
    arrow(draw, (420, 305), (470, 305))
    arrow(draw, (720, 305), (1245, 265))
    arrow(draw, (420, 330), (950, 435))
    arrow(draw, (420, 350), (950, 605))
    arrow(draw, (1370, 320), (1370, 420))
    draw.text((95, 760), "Inventory usa 2 réplicas distribuidas por hostname para cubrir el criterio multi-nodo.", fill="#111111", font=small)
    img.save(DIAGRAM)


def set_doc_styles(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    for name, size in [("Heading 1", 15), ("Heading 2", 12.5), ("Heading 3", 11)]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)


def add_doc_paragraph(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        p.add_run(text[len(bold_prefix) :])
    else:
        p.add_run(text)


def add_doc_table(doc: Document, headers: list[str], rows: list[list[str]], font_size: float = 8.5) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = text
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(font_size)
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[idx].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(font_size)


def add_doc_image(doc: Document, image_path: Path, caption: str, width: float = 6.25) -> None:
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        r.italic = True
        r.font.size = Pt(9)


def build_docx() -> None:
    build_diagram()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    set_doc_styles(doc)

    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Inches(3.35))
    for text, size, bold in [
        ("UNIVERSIDAD POLITÉCNICA SALESIANA", 15, True),
        (CAREER.upper(), 12, True),
        (TITLE.upper(), 18, True),
        (SUBTITLE, 12, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
    for text in [
        f"Integrantes: {AUTHORS}",
        f"Materia: {COURSE}",
        f"Docente: {TEACHER}",
        "Periodo académico: 2026-2027",
        f"Fecha: {DATE}",
        LOCATION,
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(text)
    doc.add_page_break()

    doc.add_heading("Resumen", level=1)
    add_doc_paragraph(
        doc,
        "Este informe documenta una práctica de tolerancia a fallas aplicada a un sistema simplificado de reservas de entradas desplegado en Kubernetes. "
        "La solución conserva seis componentes solicitados: API Gateway, Reservas Core, Inventario, Pagos, Notificaciones y Base de Datos. "
        "El laboratorio se ejecuto sobre un cluster minikube de dos nodos y distribuyo el servicio crítico de inventario con dos réplicas. "
        "Se implementaron cuatro defensas reales: retry con backoff para inventario, timeout y circuit breaker para pagos, fallback para notificaciónes "
        "y rate limiting en el gateway. Los dos fallos restantes, base de datos intermitente y condición de carrera, se analizan teóricamente con propuestas "
        "de solución de nivel producción.",
    )
    add_doc_paragraph(
        doc,
        "Palabras clave: Kubernetes, tolerancia a fallas, microservicios, resiliencia, circuit breaker, retry, fallback, rate limiting, consistencia.",
    )

    doc.add_heading("Contenido", level=1)
    add_doc_table(
        doc,
        ["Seccion", "Descripcion"],
        [
            ["1", "Objetivo, alcance y criterios de evaluación"],
            ["2", "Arquitectura y despliegue multi-nodo"],
            ["3", "Catálogo de fallos e implementación de defensas"],
            ["4", "Evidencia experimental y guion de demo"],
            ["5", "Análisis teórico de fallos no implementados"],
            ["6", "Declaración de IA y referencias bibliográficas"],
        ],
        9,
    )

    doc.add_heading("1. Objetivo, alcance y criterios", level=1)
    add_doc_paragraph(
        doc,
        "El objetivo fue construir una práctica reproducible en la que los fallos no se describan solo de forma teórica, sino que se provoquen sobre una infraestructura Kubernetes real. "
        "El alcance cubre despliegue, scripts de inyeccion, evidencias, demo y analisis de dos escenarios restantes.",
    )
    add_doc_table(
        doc,
        ["Criterio de rúbrica", "Cumplimiento en el proyecto"],
        [
            ["Cluster de al menos dos nodos", "Perfil minikube ticket-lab con nodos ticket-lab y ticket-lab-m02."],
            ["Componente crítico replicado", "inventory con 2 réplicas distribuidas por topologySpreadConstraints."],
            ["Seis componentes conservados", "gateway, reservations, inventory, payments, notifications y postgres."],
            ["Cuatro fallos implementados", "Inventory down, payments lento, notifications down y traffic spike."],
            ["Dos fallos analizados", "Base de datos intermitente y condición de carrera."],
            ["Evidencia y demo", "Capturas en docs/evidence y scripts numerados en scripts/."],
        ],
        8.3,
    )

    doc.add_heading("2. Arquitectura y despliegue", level=1)
    add_doc_table(doc, ["Componente", "Rol", "Responsabilidad", "Deployment"], COMPONENTS, 8.2)
    add_doc_image(doc, DIAGRAM, "Figura 1. Arquitectura y distribucion multi-nodo del laboratorio.", 6.4)
    add_doc_paragraph(
        doc,
        "El despliegue usa Services internos de Kubernetes para resolver dependencias por nombre DNS. Los contenedores exponen endpoints /health y los manifiestos incluyen readiness/liveness probes. "
        "El namespace ticket-lab separa los recursos de la práctica. El uso de 2 réplicas en inventory evita que todo el componente crítico quede concentrado en un único nodo.",
    )

    doc.add_heading("3. Catálogo de fallos y matriz técnica", level=1)
    add_doc_table(doc, ["Fallo", "Tipo", "Inyeccion", "Defensa", "Estado"], FAILURE_MATRIX, 7.4)

    doc.add_heading("4. Mecanismos implementados", level=1)
    for item in IMPLEMENTED_FAILURES:
        doc.add_heading(item["title"], level=2)
        add_doc_paragraph(doc, f"Riesgo tecnico: {item['risk']}", "Riesgo tecnico:")
        add_doc_paragraph(doc, f"Comando de demo: {item['command']}", "Comando de demo:")
        add_doc_paragraph(doc, f"Defensa implementada: {item['defense']}", "Defensa implementada:")
        add_doc_paragraph(doc, f"Justificacion del patrón: {item['why']}", "Justificacion del patrón:")
        add_doc_paragraph(doc, f"Resultado observado: {item['result']}", "Resultado observado:")
        add_doc_image(doc, EVIDENCE / item["evidence"], f"Evidencia: {item['title']}.", 6.15)

    doc.add_heading("5. Guion de demo en vivo", level=1)
    add_doc_table(
        doc,
        ["Tiempo", "Responsable", "Accion", "Evidencia esperada"],
        [
            ["0:00-2:00", "Alexander", "Mostrar nodos y pods.", "Dos nodos Ready e inventory replicado."],
            ["2:00-4:00", "Jean Pierre", "Ejecutar flujo base.", "HTTP 200 y reserva confirmed."],
            ["4:00-6:30", "Alexander", "Inventory down y recuperacion.", "HTTP 503 controlado."],
            ["6:30-9:00", "Jean Pierre", "Payments lento.", "Circuit breaker open."],
            ["9:00-11:30", "Alexander", "Notifications down.", "Reserva confirmed con pending."],
            ["11:30-13:30", "Jean Pierre", "Traffic spike.", "HTTP 429 en el gateway."],
            ["13:30-15:00", "Ambos", "Conclusiones y límites.", "Relación con fallos teoricos."],
        ],
        8,
    )

    doc.add_heading("6. Análisis teórico de fallos restantes", level=1)
    doc.add_heading("6.1 Base de Datos Intermitente", level=2)
    add_doc_paragraph(
        doc,
        "Este fallo ocurre cuando el servicio de inventario pierde conectividad intermitente con PostgreSQL durante escrituras. En terminos de sistemas distribuidos, una particion obliga a decidir entre disponibilidad y consistencia. "
        "En reservas, aceptar disponibilidad sin escritura durable puede producir ventas fantasma; por ello, la confirmacion debe depender de una transaccion persistida.",
    )
    add_doc_paragraph(
        doc,
        "Solucion propuesta: base de datos administrada con failover, readiness probes para retirar pods no saludables, pool de conexiones con límites, retries con backoff y jitter para errores transitorios, idempotency keys y patrón outbox para eventos posteriores.",
    )
    add_doc_paragraph(
        doc,
        "Pseudocodigo: begin transaction; select seat for update; si esta disponible, crear reserva con idempotency key y evento outbox; commit; si hay error transitorio, rollback y retry con backoff; si excede el limite, devolver error controlado.",
    )
    doc.add_heading("6.2 Condición de Carrera", level=2)
    add_doc_paragraph(
        doc,
        "La condición de carrera aparece cuando dos usuarios intentan comprar el mismo asiento al mismo tiempo. Si la verificacion y el descuento de inventario se hacen sin aislamiento, ambos procesos pueden leer disponibilidad y confirmar ventas duplicadas.",
    )
    add_doc_paragraph(
        doc,
        "Solucion propuesta: transacciones ACID, SELECT FOR UPDATE, restriccion unica por asiento activo, holds temporales con expiracion e idempotencia por solicitud. El hold se confirma solo si no expiro y el pago corresponde a la reserva.",
    )
    add_doc_paragraph(
        doc,
        "Pseudocodigo: begin transaction; select seat for update; si status != available, rollback y conflict; crear hold con expires_at; update seat=held; commit. Para confirmar: bloquear hold, validar expiracion, marcar seat=reserved y commit.",
    )

    doc.add_heading("7. Reproducibilidad", level=1)
    add_doc_table(
        doc,
        ["Paso", "Comando"],
        [
            ["Preparar cluster", ".\\scripts\\00-minikube-ensure-two-nodes.ps1"],
            ["Construir imagenes", ".\\scripts\\01-build-load-images.ps1"],
            ["Desplegar", ".\\scripts\\02-deploy-k8s.ps1"],
            ["Port-forward", ".\\scripts\\03-port-forward.ps1"],
            ["Evidencia base", ".\\scripts\\04-baseline-evidence.ps1"],
            ["Reiniciar laboratorio", ".\\scripts\\12-reset-lab.ps1"],
        ],
        8.5,
    )

    doc.add_heading("8. Declaración de uso de inteligencia artificial", level=1)
    add_doc_paragraph(
        doc,
        "Durante el desarrollo se útilizo ChatGPT 5.5 como apoyo para organizar la práctica, revisar la correspondencia con la rúbrica, proponer la estructura del informe, depurar comandos y mejorar la redacción técnica. "
        "La herramienta no sustituyo la ejecución práctica: los comandos fueron corridos sobre el entorno local, las capturas fueron generadas por los integrantes y los resultados se verificaron contra el cluster Kubernetes desplegado.",
    )
    add_doc_paragraph(
        doc,
        "Se revisaron manualmente los resultados antes de incorporarlos al informe. Las decisiones técnicas principales fueron: implementar cuatro fallos en vivo, mantener dos fallos como analisis teorico, usar inventory como componente crítico replicado y separar errores de negocio HTTP 409 de protección por sobrecarga HTTP 429.",
    )

    doc.add_heading("9. Conclusiones", level=1)
    add_doc_paragraph(
        doc,
        "La práctica demuestra que la tolerancia a fallas depende de ubicar cada patrón en el punto correcto de la arquitectura. Retry con backoff es útil cuando la dependencia puede recuperarse; timeout y circuit breaker protegen frente a latencia sostenida; fallback mantiene disponibilidad cuando una dependencia no crítica falla; y rate limiting reduce la presión antes de que el tráfico alcance servicios internos.",
    )
    add_doc_paragraph(
        doc,
        "El laboratorio cumple el objetivo de experimentar fallos en infraestructura Kubernetes multi-nodo, con evidencia reproducible y comandos claros para demo. Para un entorno productivo se requeririan observabilidad, trazas distribuidas, persistencia durable de eventos, politicas de red mas estrictas y almacenamiento de alta disponibilidad.",
    )

    doc.add_heading("Referencias bibliográficas", level=1)
    for ref in REFERENCES:
        doc.add_paragraph(ref)

    doc.add_page_break()
    doc.add_heading("Anexo A. Evidencias finales", level=1)
    for ev in EVIDENCE_LIST:
        add_doc_image(doc, EVIDENCE / ev.filename, ev.caption, 6.15)

    doc.save(DOCX_OUT)


def p(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(text.replace("&", "&amp;"), style)


def pdf_image(path: Path, max_width: float = 6.4 * inch) -> PdfImage:
    with Image.open(path) as img:
        w, h = img.size
    ratio = max_width / w
    return PdfImage(str(path), width=max_width, height=h * ratio)


def table_style(font_size: float = 8.0, header_bg: str = "#111111") -> TableStyle:
    return TableStyle(
        [
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(header_bg)),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), font_size),
            ("GRID", (0, 0), (-1, -1), 0.35, colors.black),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]
    )


def para_table(rows: list[list[str]], widths: list[float], style: ParagraphStyle, repeat_rows: int = 1, font_size: float = 8.0) -> Table:
    converted = [[p(cell, style) for cell in row] for row in rows]
    table = Table(converted, colWidths=widths, repeatRows=repeat_rows)
    table.setStyle(table_style(font_size=font_size))
    return table


def build_pdf() -> None:
    build_diagram()
    styles = getSampleStyleSheet()
    title = ParagraphStyle("Title", parent=styles["Title"], alignment=TA_CENTER, fontSize=18, leading=22, textColor=colors.black)
    subtitle = ParagraphStyle("Subtitle", parent=styles["BodyText"], alignment=TA_CENTER, fontSize=10.5, leading=14)
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontSize=14, leading=18, textColor=colors.black, spaceBefore=10, spaceAfter=6)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=11.5, leading=15, textColor=colors.black, spaceBefore=8, spaceAfter=4)
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontSize=9, leading=12, alignment=TA_LEFT, spaceAfter=5)
    small = ParagraphStyle("Small", parent=body, fontSize=7.4, leading=9)
    caption = ParagraphStyle("Caption", parent=styles["Italic"], alignment=TA_CENTER, fontSize=8, leading=10)

    story: list = []
    if LOGO.exists():
        story.append(PdfImage(str(LOGO), width=3.2 * inch, height=0.88 * inch))
    story.extend(
        [
            Spacer(1, 0.25 * inch),
            p("UNIVERSIDAD POLITÉCNICA SALESIANA", title),
            p(CAREER.upper(), subtitle),
            Spacer(1, 0.2 * inch),
            p(TITLE.upper(), title),
            p(SUBTITLE, subtitle),
            Spacer(1, 0.2 * inch),
            p(f"Integrantes: {AUTHORS}", subtitle),
            p(f"Materia: {COURSE}", subtitle),
            p(f"Docente: {TEACHER}", subtitle),
            p("Periodo académico: 2026-2027", subtitle),
            p(f"Fecha: {DATE}", subtitle),
            p(LOCATION, subtitle),
            PageBreak(),
            p("Resumen", h1),
            p(
                "Este informe documenta una práctica de tolerancia a fallas aplicada a un sistema simplificado de reservas de entradas desplegado en Kubernetes. "
                "La solución conserva seis componentes solicitados y ejecuta cuatro fallos reales sobre un cluster de dos nodos: inventario caído, pasarela lenta, notificaciónes caídas y diluvio de peticiones. "
                "Los dos fallos restantes se analizan teóricamente con propuestas de solución de nivel producción.",
                body,
            ),
            p("Palabras clave: Kubernetes, tolerancia a fallas, microservicios, resiliencia, circuit breaker, retry, fallback, rate limiting, consistencia.", body),
            p("1. Objetivo, alcance y criterios", h1),
            p(
                "El objetivo fue construir una práctica reproducible en la que los fallos se provoquen sobre infraestructura Kubernetes real. El alcance cubre despliegue, scripts de inyeccion, evidencias, demo y analisis de los escenarios restantes.",
                body,
            ),
        ]
    )

    rubric_rows = [
        ["Criterio", "Cumplimiento"],
        ["Cluster de al menos dos nodos", "Perfil minikube ticket-lab con nodos ticket-lab y ticket-lab-m02."],
        ["Componente crítico replicado", "inventory con 2 réplicas distribuidas por topologySpreadConstraints."],
        ["Seis componentes", "gateway, reservations, inventory, payments, notifications y postgres."],
        ["Cuatro fallos implementados", "Inventory down, payments lento, notifications down y traffic spike."],
        ["Dos fallos analizados", "Base de datos intermitente y condición de carrera."],
        ["Evidencia y demo", "Capturas en docs/evidence y scripts numerados en scripts/."],
    ]
    story.append(para_table(rubric_rows, [2.1 * inch, 4.4 * inch], body, font_size=7.8))

    story.extend([p("2. Arquitectura y despliegue", h1)])
    comp_rows = [["Componente", "Rol", "Responsabilidad", "Deployment"]] + COMPONENTS
    story.append(para_table(comp_rows, [1.15 * inch, 1.2 * inch, 2.95 * inch, 1.05 * inch], small, font_size=6.8))
    story.append(Spacer(1, 0.12 * inch))
    story.append(pdf_image(DIAGRAM))
    story.append(p("Figura 1. Arquitectura y distribucion multi-nodo del laboratorio.", caption))
    story.append(
        p(
            "El despliegue usa Services internos de Kubernetes, endpoints /health y readiness/liveness probes. Inventory se replica en dos nodos para evitar concentrar todo el componente crítico en una sola maquina.",
            body,
        )
    )

    story.extend([p("3. Catálogo de fallos y matriz técnica", h1)])
    failure_rows = [["Fallo", "Tipo", "Inyeccion", "Defensa", "Estado"]] + FAILURE_MATRIX
    story.append(para_table(failure_rows, [1.25 * inch, 0.8 * inch, 1.7 * inch, 1.65 * inch, 1.05 * inch], small, font_size=6.6))

    story.append(PageBreak())
    story.append(p("4. Mecanismos implementados", h1))
    for item in IMPLEMENTED_FAILURES:
        story.extend(
            [
                p(item["title"], h2),
                p(f"<b>Riesgo tecnico:</b> {item['risk']}", body),
                p(f"<b>Comando:</b> {item['command']}", body),
                p(f"<b>Defensa:</b> {item['defense']}", body),
                p(f"<b>Justificacion:</b> {item['why']}", body),
                p(f"<b>Resultado:</b> {item['result']}", body),
            ]
        )

    story.append(p("5. Guion de demo en vivo", h1))
    demo_rows = [
        ["Tiempo", "Responsable", "Accion", "Evidencia esperada"],
        ["0:00-2:00", "Alexander", "Mostrar nodos y pods.", "Dos nodos Ready e inventory replicado."],
        ["2:00-4:00", "Jean Pierre", "Ejecutar flujo base.", "HTTP 200 y reserva confirmed."],
        ["4:00-6:30", "Alexander", "Inventory down y recuperacion.", "HTTP 503 controlado."],
        ["6:30-9:00", "Jean Pierre", "Payments lento.", "Circuit breaker open."],
        ["9:00-11:30", "Alexander", "Notifications down.", "Reserva confirmed con pending."],
        ["11:30-13:30", "Jean Pierre", "Traffic spike.", "HTTP 429 en gateway."],
        ["13:30-15:00", "Ambos", "Conclusiones y límites.", "Relación con fallos teoricos."],
    ]
    story.append(para_table(demo_rows, [0.9 * inch, 1.05 * inch, 2.35 * inch, 2.1 * inch], small, font_size=6.8))

    story.append(PageBreak())
    story.extend(
        [
            p("6. Análisis teórico de fallos restantes", h1),
            p("6.1 Base de Datos Intermitente", h2),
            p(
                "Este fallo ocurre cuando inventory pierde conectividad intermitente con PostgreSQL durante escrituras. Desde CAP, una particion obliga a decidir entre disponibilidad y consistencia; en reservas, confirmar sin escritura durable puede crear ventas fantasma.",
                body,
            ),
            p(
                "Solucion propuesta: base de datos con failover, readiness probes, pool de conexiones, retries con backoff y jitter, idempotency keys y patrón outbox.",
                body,
            ),
            p(
                "<b>Pseudocodigo:</b> begin transaction; select seat for update; si esta disponible, crear reserva con idempotency key y evento outbox; commit; ante error transitorio, rollback y retry; si excede el limite, error controlado.",
                body,
            ),
            p("6.2 Condición de Carrera", h2),
            p(
                "La condición de carrera ocurre cuando dos usuarios intentan comprar el mismo asiento al mismo tiempo. Sin aislamiento transaccional, ambos procesos pueden observar disponibilidad y confirmar ventas duplicadas.",
                body,
            ),
            p(
                "Solucion propuesta: transacciones ACID, SELECT FOR UPDATE, restriccion unica por asiento activo, holds temporales con expiracion e idempotencia por solicitud.",
                body,
            ),
            p(
                "<b>Pseudocodigo:</b> begin transaction; select seat for update; si status != available, rollback y conflict; crear hold con expires_at; update seat=held; commit. Para confirmar, bloquear hold, validar expiracion y marcar seat=reserved.",
                body,
            ),
            p("7. Reproducibilidad", h1),
        ]
    )
    repro_rows = [
        ["Paso", "Comando"],
        ["Preparar cluster", ".\\scripts\\00-minikube-ensure-two-nodes.ps1"],
        ["Construir imagenes", ".\\scripts\\01-build-load-images.ps1"],
        ["Desplegar", ".\\scripts\\02-deploy-k8s.ps1"],
        ["Port-forward", ".\\scripts\\03-port-forward.ps1"],
        ["Evidencia base", ".\\scripts\\04-baseline-evidence.ps1"],
        ["Reiniciar laboratorio", ".\\scripts\\12-reset-lab.ps1"],
    ]
    story.append(para_table(repro_rows, [1.8 * inch, 4.7 * inch], body, font_size=7.8))
    story.extend(
        [
            p("8. Declaración de uso de inteligencia artificial", h1),
            p(
                "Durante el desarrollo se útilizo ChatGPT 5.5 como apoyo para organizar la práctica, revisar la correspondencia con la rúbrica, proponer la estructura del informe, depurar comandos y mejorar la redacción técnica. La herramienta no sustituyo la ejecución práctica: los comandos fueron corridos sobre el entorno local, las capturas fueron generadas por los integrantes y los resultados se verificaron contra el cluster Kubernetes desplegado.",
                body,
            ),
            p(
                "Se revisaron manualmente los resultados antes de incorporarlos al informe. Las decisiones técnicas principales fueron implementar cuatro fallos en vivo, dejar dos como analisis teorico, usar inventory como componente crítico replicado y separar errores de negocio HTTP 409 de protección por sobrecarga HTTP 429.",
                body,
            ),
            p("9. Conclusiones", h1),
            p(
                "La práctica demuestra que cada patrón debe ubicarse en el punto correcto: retry con backoff para dependencias transitorias, timeout y circuit breaker frente a latencia sostenida, fallback para dependencias no críticas y rate limiting en el borde del sistema.",
                body,
            ),
            p("Referencias bibliográficas", h1),
        ]
    )
    for ref in REFERENCES:
        story.append(p(ref, body))

    story.append(PageBreak())
    story.append(p("Anexo A. Evidencias finales", h1))
    for ev in EVIDENCE_LIST:
        path = EVIDENCE / ev.filename
        if path.exists():
            story.append(KeepTogether([pdf_image(path), Spacer(1, 0.04 * inch), p(ev.caption, caption)]))
            story.append(Spacer(1, 0.1 * inch))

    doc = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=letter,
        leftMargin=0.62 * inch,
        rightMargin=0.62 * inch,
        topMargin=0.65 * inch,
        bottomMargin=0.65 * inch,
    )
    doc.build(story)


def build_markdown() -> None:
    lines = [
        f"# {TITLE}",
        "",
        f"**{SUBTITLE}**",
        "",
        f"**Integrantes:** {AUTHORS}",
        f"**Repositorio:** {REPO_URL}",
        "",
        "## Resumen",
        "",
        "Este informe documenta una práctica de tolerancia a fallas aplicada a un sistema simplificado de reservas de entradas desplegado en Kubernetes. La solución conserva los seis componentes solicitados y ejecuta cuatro fallos reales sobre un cluster de dos nodos.",
        "",
        "## Cumplimiento de rúbrica",
        "",
        "- Cluster multi-nodo: perfil minikube `ticket-lab` con `ticket-lab` y `ticket-lab-m02`.",
        "- Componente crítico replicado: `inventory` con 2 réplicas distribuidas por nodo.",
        "- Cuatro fallos implementados: inventory down, payments lento, notifications down y traffic spike.",
        "- Dos fallos analizados: base de datos intermitente y condición de carrera.",
        "",
        "## Declaración IA",
        "",
        "Se útilizo ChatGPT 5.5 como apoyo para organizacion, revision de rúbrica, depuracion de comandos y redacción técnica. La ejecución práctica y las capturas fueron realizadas por los integrantes.",
        "",
        "## Referencias",
        "",
    ]
    lines.extend([f"- {ref}" for ref in REFERENCES])
    MD_OUT.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    build_docx()
    build_pdf()
    build_markdown()
    print(DOCX_OUT)
    print(PDF_OUT)
    print(MD_OUT)


if __name__ == "__main__":
    main()

